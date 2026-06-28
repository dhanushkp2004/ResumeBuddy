import io
import pdfplumber
import docx
from fastapi import HTTPException, status

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extracts text from PDF bytes using pdfplumber.
    Throws HTTPException if text cannot be extracted or if pdf is corrupt.
    """
    try:
        text_content = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            if not pdf.pages:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail="The uploaded PDF contains no pages."
                )
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
                
        extracted_text = "\n".join(text_content).strip()
        if not extracted_text:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="Could not extract any readable text from this PDF. It might be scanned or image-only."
            )
        return extracted_text
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Corrupt or invalid PDF file. Error: {str(e)}"
        )

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts text from DOCX bytes using python-docx.
    """
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        text_content = []
        
        # Read normal paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
                
        # Read text from tables (common in resumes)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in row_text: # avoid duplicating cell text in merged cells
                        row_text.append(cell_text)
                if row_text:
                    text_content.append(" | ".join(row_text))
                    
        extracted_text = "\n".join(text_content).strip()
        if not extracted_text:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="Could not extract any readable text from this DOCX file."
            )
        return extracted_text
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Corrupt or invalid DOCX file. Error: {str(e)}"
        )

def parse_resume(filename: str, file_bytes: bytes) -> str:
    """
    Wrapper function to check extension and parse accordingly.
    """
    # Check file size (10 MB limit = 10 * 1024 * 1024 bytes)
    max_size = 10 * 1024 * 1024
    if len(file_bytes) > max_size:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail="File size exceeds the maximum limit of 10 MB."
        )

    ext = filename.split(".")[-1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext == "docx":
        return extract_text_from_docx(file_bytes)
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file format '.{ext}'. Only PDF and DOCX files are allowed."
        )
